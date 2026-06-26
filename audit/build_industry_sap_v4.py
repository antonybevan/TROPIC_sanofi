from __future__ import annotations

import csv
import re
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "TROPIC_SAP_v4.0_industry_grade.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
RISK = RGBColor(156, 0, 6)
AMBER = RGBColor(156, 101, 0)
GREEN = RGBColor(0, 97, 0)
WHITE = RGBColor(255, 255, 255)


def ascii_clean(s: str) -> str:
    repl = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2011": "-",
        "\u2212": "-",
        "\u2265": ">=",
        "\u2264": "<=",
        "\u03b1": "alpha",
        "\u03bc": "u",
        "\u00b2": "^2",
        "\u00b3": "^3",
        "\u00b1": "+/-",
        "\u00b7": ".",
        "\u00a7": "Section ",
        "\u00d7": "x",
        "\u03c4": "tau",
        "\u2020": "*",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
    }
    for a, b in repl.items():
        s = s.replace(a, b)
    return s


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 8.3) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(ascii_clean(str(text)))
    set_run_font(r, size=size, color=color or DARK, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Heading 1", 15.5, 13, 6),
        ("Heading 2", 12.5, 9, 4),
        ("Heading 3", 10.8, 6, 2),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = BLUE
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(9.5)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.line_spacing = 1.05


def set_header_footer(doc: Document) -> None:
    for sec in doc.sections:
        header = sec.header
        p = header.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("TROPIC / EFC6193 SAP v4.0 - Controlled Draft")
        set_run_font(r, size=8, color=MUTED)
        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = fp.add_run("Confidential clinical programming work product - approval required before submission use")
        set_run_font(rr, size=8, color=MUTED)


def add_p(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    text = ascii_clean(text)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(ascii_clean(item))


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Evidence source: ")
    set_run_font(r, size=7.6, bold=True, color=MUTED)
    rr = p.add_run(ascii_clean(text))
    set_run_font(rr, size=7.6, color=MUTED)


def add_callout(doc: Document, text: str, kind: str = "note") -> None:
    fills = {"note": "EAF3F8", "risk": "FCE4D6", "ok": "E2F0D9", "warn": "FFF2CC"}
    colors = {"note": BLUE, "risk": RISK, "ok": GREEN, "warn": AMBER}
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_shading(c, fills.get(kind, "EAF3F8"))
    cell_text(c, text, bold=True, color=colors.get(kind, BLUE), size=8.8)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: float = 8.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, color=WHITE, size=font_size)
        set_cell_shading(t.rows[0].cells[i], "1F4E79")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val, size=font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths[: len(row.cells)]):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(ascii_clean(text), level=level)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def rows_from_workbook(sheet: str) -> list[dict[str, str]]:
    wb = load_workbook(ROOT / "00_specifications" / "ADaM_spec.xlsx", data_only=True, read_only=True)
    ws = wb[sheet]
    values = list(ws.iter_rows(values_only=True))
    headers = [str(x) if x is not None else "" for x in values[0]]
    out = []
    for row in values[1:]:
        if any(x is not None for x in row):
            out.append({headers[i]: "" if row[i] is None else str(row[i]) for i in range(min(len(headers), len(row)))})
    return out


def sap_v3_tfl_rows() -> list[dict[str, str]]:
    doc = Document(ROOT / "TROPIC_SAP_v3.0.docx")
    t = doc.tables[20]
    headers = [c.text.strip() for c in t.rows[0].cells]
    rows = []
    produced = {p.name for p in (ROOT / "09_tfl" / "output").glob("**/*") if p.is_file()}
    produced_ids = {re.match(r"([FTL]-\d+(?:-\d+)?)", p).group(1) for p in produced if re.match(r"([FTL]-\d+(?:-\d+)?)", p)}
    produced_ids.update({"T-11", "T-17", "T-20", "T-21"})
    for row in t.rows[1:]:
        vals = [c.text.replace("\n", " ").strip() for c in row.cells]
        d = dict(zip(headers, vals))
        oid = d.get("ID", "")
        if oid.startswith("L-01"):
            status = "Exclude current file; regenerate from source after shell approval."
        elif oid in produced_ids or oid.split("-")[0] + "-" + oid.split("-")[1] in produced_ids:
            status = "Current physical output exists; requires v4 shell/QC status check."
        elif oid.startswith("T-14") or oid.startswith("T-13") or oid.startswith("T-15") or oid.startswith("T-16") or oid in {"F-12-2"}:
            status = "Planned; not currently produced as a controlled output."
        else:
            status = "Target output; implementation status requires reconciliation."
        d["Status"] = status
        rows.append(d)
    return rows


def sap_v3_table(index_1_based: int) -> tuple[list[str], list[list[str]]]:
    doc = Document(ROOT / "TROPIC_SAP_v3.0.docx")
    t = doc.tables[index_1_based - 1]
    headers = [ascii_clean(c.text.replace("\n", " ").strip()) for c in t.rows[0].cells]
    rows = []
    for row in t.rows[1:]:
        vals = [ascii_clean(c.text.replace("\n", " ").strip()) for c in row.cells]
        if any(vals):
            rows.append(vals)
    return headers, rows


def findings_rows(severity: str | None = None) -> list[list[str]]:
    rows = []
    p = ROOT / "audit" / "findings_register.csv"
    if p.exists():
        with p.open(newline="") as f:
            for r in csv.DictReader(f):
                if severity and r["severity"] != severity:
                    continue
                rows.append([r["ID"], r["severity"], r["category"], r["remediation"]])
    return rows


def source_register_rows() -> list[list[str]]:
    return [
        [
            "Protocol Amendment 5, 21-Jul-2008",
            "01_raw_source/Sanofi Study Protocol Tropic.pdf",
            "Objectives, treatment arms, randomization, strata, populations, endpoint definitions, sample size, CTCAE grading, assessment schedule.",
            "High for intended trial design; source is sponsor protocol PDF available in repo.",
        ],
        [
            "de Bono et al. Lancet 2010 corrected publication",
            "01_raw_source/reference_literature/de_bono_lancet_2010.pdf",
            "Final cut-off, published ITT/safety counts, primary/secondary results, safety tables, references to CTCAE v3.0 and MedDRA v12.0.",
            "High for public reconciliation targets; not a substitute for original locked analysis datasets.",
        ],
        [
            "Sanofi CRF TROPIC",
            "01_raw_source/Sanofi CRF Tropic.pdf",
            "CRF capture context for eligibility, RECIST, labs, pain, AE, exposure, concomitant medication and follow-up forms.",
            "Medium-high for collection design; not sufficient alone for derivation algorithms.",
        ],
        [
            "PDS / Sanofi public SDTM release",
            "01_raw_source/real_sdtm/*.sas7bdat",
            "Real patient-level MP arm data in the current repo pipeline.",
            "High for MP arm analyses; incomplete for original two-arm confirmatory inference.",
        ],
        [
            "ADaM specification workbook",
            "00_specifications/ADaM_spec.xlsx",
            "ADaM datasets, variables, codelists, methods, value-level metadata and WhereClauses.",
            "Implementation baseline; currently requires traceability completion before release authority.",
        ],
        [
            "Audit evidence pack",
            "audit/AUDIT_REPORT.md and audit/findings_register.csv",
            "Critical and Major defects that must be closed or risk-accepted before submission-readiness claims.",
            "High for remediation control; not a clinical source document.",
        ],
        [
            "Reconstruction programs and digitized KM inputs",
            "01_raw_source/reconstruct_cbzp_guyot.R; 01_raw_source/reconstruct_cbzp_arm.R; 01_raw_source/guyot_digitised/*",
            "Synthetic/reconstructed CbzP data generation for demonstration.",
            "Valid only for explicitly non-confirmatory demonstration unless formally justified and approved.",
        ],
    ]


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STATISTICAL ANALYSIS PLAN")
    set_run_font(r, size=24, color=BLUE, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("TROPIC / EFC6193 - Industry-Grade Clinical Programming Baseline")
    set_run_font(r2, size=15, color=DARK, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Cabazitaxel + Prednisone vs Mitoxantrone + Prednisone in Metastatic Castration-Resistant Prostate Cancer")
    set_run_font(r3, size=11.5, color=MUTED, italic=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["Document attribute", "Value"],
        [
            ["Version", "4.0 controlled draft"],
            ["Date", date.today().isoformat()],
            ["Status", "Sponsor-approvable SAP draft for remediation and programming control; not sponsor-approved until signed."],
            ["Study", "EFC6193 / TROPIC / NCT00417079"],
            ["Primary data limitation", "Current repo contains real MP patient-level data and reconstructed/synthetic CbzP data. Full two-arm confirmatory inference requires complete authoritative two-arm patient-level data."],
            ["Document purpose", "Define the statistical and programming authority needed to rebuild the pipeline as a serious 2026 clinical programming work product."],
        ],
        widths=[1.75, 4.75],
        font_size=8.5,
    )
    add_callout(
        doc,
        "Control statement: This SAP is intentionally stricter than the current repo. Programs, TFLs, Define-XML, reviewer guides and eCTD outputs must conform to this SAP and its release gates. Generated artifacts cannot retroactively establish prespecification.",
        "risk",
    )
    add_table(
        doc,
        ["Approval role", "Name / signature", "Date", "Decision"],
        [
            ["Lead Statistician", "", "", "Approve / Revise / Reject"],
            ["Lead Statistical Programmer", "", "", "Approve / Revise / Reject"],
            ["Independent QC Programmer", "", "", "Approve / Revise / Reject"],
            ["Clinical Domain Lead", "", "", "Approve / Revise / Reject"],
            ["Data Standards Lead", "", "", "Approve / Revise / Reject"],
            ["Quality / Validation Representative", "", "", "Approve / Revise / Reject"],
        ],
        widths=[1.6, 2.15, 1.0, 1.75],
        font_size=8.5,
    )
    page_break(doc)


def add_toc(doc: Document) -> None:
    add_heading(doc, "Table of contents", 1)
    sections = [
        "1. Administrative control and document status",
        "2. Evidence authority and source hierarchy",
        "3. Study overview",
        "4. Data provenance and analysis layers",
        "5. Analysis populations",
        "6. General statistical principles",
        "7. Estimand framework",
        "8. Multiplicity and Type I error control",
        "9. Primary efficacy analysis - overall survival",
        "10. Secondary efficacy analyses",
        "11. Safety analyses",
        "12. Exposure, dose modification and Project Optimus exploratory analyses",
        "13. Subgroup, sensitivity and supplementary analyses",
        "14. Data handling, missing data, partial dates and visit windows",
        "15. CDISC implementation and metadata traceability",
        "16. Programming, independent QC and release gates",
        "17. TFL catalog and shell control",
        "18. Known limitations and remediation controls",
        "19. References",
        "Appendix A. Source register",
        "Appendix B. Endpoint algorithm details",
        "Appendix C. ADaM dataset and metadata requirements",
        "Appendix D. Full TFL catalog",
        "Appendix E. Published validation targets",
        "Appendix F. Safety, death and exposure publication targets",
        "Appendix G. Subgroup target register",
        "Appendix H. Visit-window and protocol-correction registers",
        "Appendix I. Full ADaM variable specification summary",
        "Appendix J. Critical audit closure checklist",
    ]
    add_bullets(doc, sections)
    page_break(doc)


def build_doc() -> Document:
    doc = Document()
    style_doc(doc)
    set_header_footer(doc)
    add_title_page(doc)
    add_toc(doc)

    add_heading(doc, "1. Administrative control and document status", 1)
    add_p(
        doc,
        "This document is the controlling Statistical Analysis Plan (SAP) draft for rebuilding the TROPIC clinical programming pipeline as a serious regulated-workflow artifact. It is not the original Sanofi SAP and does not assert sponsor approval. It is designed to be approvable after sponsor/statistical review, source-data reconciliation, and closure of release gates.",
    )
    add_table(
        doc,
        ["Control item", "Requirement"],
        [
            ["Change control", "Any change after approval requires a versioned SAP amendment, rationale, affected outputs, affected ADaM/Define metadata, and signer approval."],
            ["Traceability", "Every analysis result must trace to SAP section, TFL shell, ADaM dataset/variable(s), SDTM/source predecessor(s), program, QC program and run manifest."],
            ["Evidence standard", "No claim is accepted without a file, page/section, dataset, program or validator report reference."],
            ["Submission readiness", "The current repository is not submission-ready until audit Critical findings F-001 through F-005 are closed and Major findings are either closed or formally risk-accepted."],
            ["Portfolio constraint", "The repository may remain public/portfolio-oriented, but clinical claims must be controlled as non-confirmatory unless supported by authoritative complete data and approval."],
        ],
        widths=[1.5, 5.0],
    )
    add_heading(doc, "1.1 Roles and responsibilities", 2)
    add_table(
        doc,
        ["Role", "Accountability"],
        [
            ["Lead Statistician", "Owns SAP scientific content, estimands, multiplicity, statistical models, interpretation and final approval."],
            ["Lead Statistical Programmer", "Owns production programming implementation, ADaM/TFL standards alignment, run manifest and programmer sign-off."],
            ["Independent QC Programmer", "Owns independent derivation and result reproduction, discrepancy triage and QC sign-off."],
            ["Clinical Domain Lead", "Confirms endpoint clinical meaning, event adjudication assumptions, safety groupings and medically important caveats."],
            ["Data Standards Lead", "Owns CDISC conformance, Define-XML/ARM, controlled terminology and metadata/data concordance."],
            ["Quality / Validation Representative", "Owns process validation evidence, Part 11 control assessment, audit-trail retention and release readiness."],
        ],
        widths=[1.55, 4.95],
    )
    add_heading(doc, "1.2 Document deliverable set controlled by this SAP", 2)
    add_bullets(
        doc,
        [
            "ADaM production datasets and transport files.",
            "Independent R/pharmaverse validation datasets and reconciliation outputs.",
            "Tables, figures and listings in the controlled TFL catalog.",
            "Define-XML v2.1 with ARM and value-level metadata.",
            "ADRG, SDRG, traceability matrix, conformance reports and eCTD study-data package.",
            "Run manifest tying source data, programs, logs, outputs, validator results and hashes to one reproducible execution.",
        ],
    )

    add_heading(doc, "2. Evidence authority and source hierarchy", 1)
    add_p(
        doc,
        "The SAP resolves conflicts using an explicit source hierarchy. Protocol and publication evidence are clinical/statistical sources; generated code and documents are implementation evidence only.",
    )
    add_table(
        doc,
        ["Rank", "Source authority", "Governance use"],
        [
            ["1", "Protocol Amendment 5, 21-Jul-2008", "Trial objectives, design, populations, endpoint concepts, randomization strata, treatment schedules, sample size assumptions and planned primary analysis."],
            ["2", "Corrected de Bono et al. Lancet 2010 publication", "Public final-analysis cut-off, published ITT/safety counts, efficacy/safety results, and trial-era terminology confirmation."],
            ["3", "Patient-level source data made available through PDS/Sanofi", "Actual data source for repo-executable real MP analyses and source timing limitations."],
            ["4", "CRF and source define metadata", "Collection context and source-domain interpretation."],
            ["5", "This SAP v4.0 after approval", "Programming and analysis control for this project."],
            ["6", "ADaM spec, Define-XML, reviewer guides, programs and TFLs", "Implementation artifacts; must conform to approved SAP and source evidence."],
        ],
        widths=[0.5, 2.2, 3.8],
    )
    add_source(
        doc,
        "Protocol PDF pages 3, 10-12; Lancet PDF pages 1, 4-5, 8; audit/AUDIT_REPORT.md; official FDA/CDISC/ICH sources listed in References.",
    )

    add_heading(doc, "3. Study overview", 1)
    add_table(
        doc,
        ["Item", "Specification"],
        [
            ["Study identifier", "EFC6193 / TROPIC / NCT00417079"],
            ["Design", "Phase III, randomized, open-label, multi-center, multinational trial."],
            ["Indication", "Metastatic castration-resistant prostate cancer after docetaxel-containing therapy."],
            ["Treatment arms", "MP: mitoxantrone 12 mg/m^2 IV Day 1 every 21 days plus prednisone 10 mg daily. CbzP: cabazitaxel 25 mg/m^2 IV Day 1 every 21 days plus prednisone 10 mg daily."],
            ["Randomization", "1:1 IVRS randomization stratified by disease measurability per RECIST and ECOG performance status (0-1 versus 2)."],
            ["ITT randomized", "Published ITT: 755 total, 377 MP and 378 CbzP."],
            ["Safety treated", "Published safety: 371 per arm."],
            ["Final public cut-off", "25-Sep-2009."],
            ["Primary endpoint", "Overall survival."],
            ["Secondary endpoints", "PFS, PSA response/progression, ORR, pain response/progression, safety and tolerability."],
        ],
        widths=[1.35, 5.15],
    )
    add_heading(doc, "3.1 Treatment regimens and key administration rules", 2)
    add_table(
        doc,
        ["Arm", "Regimen", "Operational programming notes"],
        [
            ["MP", "Mitoxantrone 12 mg/m^2 IV over 15-30 minutes on Day 1 of each 21-day cycle plus prednisone 10 mg orally daily.", "Maximum 10 cycles. Dose reductions and delays summarized from EX/SUPPEX. Safety summaries use actual treatment received."],
            ["CbzP", "Cabazitaxel 25 mg/m^2 IV over 1 hour on Day 1 of each 21-day cycle plus prednisone 10 mg orally daily.", "Premedication with antihistamine, corticosteroid and H2 antagonist. In current repo, CbzP patient-level records are reconstructed/synthetic unless complete source IPD is supplied."],
        ],
        widths=[0.8, 3.0, 2.7],
    )
    add_heading(doc, "3.2 Assessment schedule", 2)
    add_table(
        doc,
        ["Assessment area", "Schedule / source basis", "Analysis implication"],
        [
            ["Survival", "Follow-up through final analysis cut-off and last known alive/death information.", "OS censoring and event status must retain source date and cut-off logic."],
            ["Tumour imaging", "Baseline, end of even-numbered treatment cycles, suspected progression and end of treatment.", "RECIST v1.0 response and TTUMOR/PFS component dates."],
            ["PSA", "Baseline and repeated on study per protocol schedule.", "Response requires confirmation; progression uses nadir-based rule by responder status."],
            ["Pain", "Daily PPI and analgesic use over 7 days before visit.", "Visit evaluability requires at least 5 of 7 daily values."],
            ["Haematology/labs", "CBC on Cycle Day 1, Day 8 and Day 15 with protocol safety monitoring.", "ADLB visit windows must include CxD1 pre-dose and nadir/recovery windows."],
            ["AE/safety", "AE collection and grading through treatment and safety follow-up.", "TEAE window and CTCAE v3.0 grading drive ADAE/ADLB outputs."],
        ],
        widths=[1.25, 2.75, 2.5],
    )
    add_source(doc, "Protocol synopsis and statistical considerations; de Bono et al. Lancet 2010 summary and methods.")

    add_heading(doc, "4. Data provenance and analysis layers", 1)
    add_callout(
        doc,
        "This is the central integrity rule for the project: real MP patient-level analyses and reconstructed CbzP demonstration analyses must never be blurred. The SAP permits both layers, but assigns different evidentiary status.",
        "warn",
    )
    add_table(
        doc,
        ["Layer", "Data basis", "Permitted claims", "Programming controls"],
        [
            ["Trial target layer", "Original randomized TROPIC design and published results.", "Defines intended clinical questions and validation targets.", "Use for objective, endpoint and validation-target specification."],
            ["Repo real-data layer", "Real MP patient-level SDTM/ADaM available in repo/PDS source.", "Permits real single-arm MP summaries and validation of MP-side derivations.", "SAS/R reconciliation must compare MP records without lossy normalization."],
            ["Synthetic comparator layer", "CbzP reconstructed from published KM curves/at-risk tables, PH scaling and fixed-seed marginal sampling.", "Permits demonstration of pipeline mechanics, shells, figures and exploratory methods only.", "Every output must label synthetic comparator and non-confirmatory status."],
            ["Future submission layer", "Complete authoritative two-arm IPD, approved SAP, clean CDISC package and validated environment.", "Permits submission-facing clinical analysis only after gates pass.", "Requires full rerun, P21/CORE/Define validation, eCTD integrity and Part 11/process evidence."],
        ],
        widths=[1.25, 2.0, 1.75, 1.5],
    )

    add_heading(doc, "5. Analysis populations", 1)
    add_table(
        doc,
        ["Population", "Definition", "Treatment assignment", "Use", "Target / current status"],
        [
            ["Intent-to-treat (ITT)", "All randomized patients.", "Randomized treatment.", "Primary efficacy and most secondary efficacy analyses.", "Published N=755. Current repo combined N=749 is not the original ITT and must not be labeled trial ITT."],
            ["Safety / all-treated", "All patients receiving at least part of one dose.", "Treatment actually received.", "Safety, exposure, dose modification, deaths within safety window.", "Published N=371 per arm. Current real layer has MP N=371."],
            ["Per-protocol", "ITT patients without major protocol deviations.", "Randomized treatment unless deviation requires exclusion.", "Sensitivity/supportive only.", "Must be source-flagged or derivable from protocol deviations."],
            ["PSA response evaluable", "Baseline PSA >=20 ug/L plus evaluable confirmatory PSA assessments.", "Randomized treatment.", "PSA response rate.", "Published 325 MP, 329 CbzP; programming must not use all PSARESP records."],
            ["Measurable disease", "Baseline measurable lesion per RECIST v1.0.", "Randomized treatment.", "ORR and tumour-response summaries.", "Published 204 MP, 201 CbzP."],
            ["Pain response evaluable", "Baseline median PPI >=2 or mean analgesic score >=10 with diary evaluability.", "Randomized treatment.", "Pain response.", "Published 168 MP, 174 CbzP."],
        ],
        widths=[1.05, 1.65, 1.25, 1.25, 1.3],
        font_size=7.8,
    )

    add_heading(doc, "6. General statistical principles", 1)
    add_heading(doc, "6.1 Analysis conventions", 2)
    add_bullets(
        doc,
        [
            "All efficacy tests are two-sided unless explicitly stated otherwise.",
            "Time-to-event analyses use days as the stored analysis scale. Months are reported as days divided by 30.4375 unless a shell specifies otherwise.",
            "Confidence intervals are two-sided 95% intervals unless the output shell specifies another interval.",
            "Counts and percentages use the relevant analysis population denominator. Percentages are displayed to one decimal place unless counts are small or shell-specific precision is required.",
            "Treatment labels are MP and CbzP. Current repo outputs must annotate whether CbzP is reconstructed/synthetic.",
            "No data-driven changes to endpoint definitions, windows, populations or estimators are permitted after SAP approval without amendment.",
        ],
    )
    add_heading(doc, "6.2 Baseline, visits and windows", 2)
    add_bullets(
        doc,
        [
            "Baseline is the last non-missing value on or before first dose unless the endpoint definition requires randomization-date baseline.",
            "Post-baseline excludes baseline and pre-dose values unless explicitly part of a Cycle Day 1 pre-dose laboratory window.",
            "CBC windows include scheduled Cycle Day 1 pre-dose, Day 8 and Day 15 assessments.",
            "When multiple values fall in a window, select one analysis record by minimum absolute distance to target day; ties use worst toxicity grade, then latest available record, unless a parameter-specific rule overrides this.",
            "Unscheduled assessments may be used for event detection if clinically valid and source-date reliability is adequate.",
        ],
    )
    add_heading(doc, "6.3 Software and reproducibility", 2)
    add_bullets(
        doc,
        [
            "Production and QC are dual-language where feasible: SAS production and R/pharmaverse validation.",
            "Each run must bind Git commit, program versions, inputs, outputs, logs, hashes, validator versions and reviewer disposition in a run manifest.",
            "Random generation for synthetic comparator data requires deterministic seed capture, seed rationale and hash-bound output.",
            "SAS and R logs must be reviewed and fail on unapproved errors, warnings, invalid INPUT, uninitialized variables and problematic merge notes.",
        ],
    )
    add_heading(doc, "6.4 Statistical model and display standards", 2)
    add_table(
        doc,
        ["Analysis type", "Primary method", "Display requirement", "QC requirement"],
        [
            ["Time-to-event", "Kaplan-Meier, stratified log-rank, stratified Cox PH where applicable.", "Median, 95% CI, events/N, censoring count, HR, CI, p-value and at-risk table for figures.", "Independent reproduction of event/censor counts, median, HR and p-value."],
            ["Binary response", "Exact/binomial CI and Fisher or CMH test if stratified comparison is justified.", "n/N, percent, 95% CI, difference/ratio if shell-approved.", "Population denominator and confirmation records independently verified."],
            ["Continuous/lab", "Descriptive statistics and shift tables; no inferential testing unless prespecified.", "N, mean, SD, median, Q1, Q3, min, max; baseline/worst shifts.", "Source units, grade derivation and selected analysis record independently verified."],
            ["Safety incidence", "Subject-level incidence by treatment, SOC/PT and grade/seriousness/action.", "n, %, denominator footnote, coding dictionary and grade version.", "Worst-grade and one-subject-once-per-category logic independently verified."],
            ["Exploratory exposure-response", "Descriptive regression/plots with clearly labeled non-confirmatory status.", "Effect estimates only if model assumptions and synthetic status are disclosed.", "Seed, derivation, model code and synthetic flags independently checked."],
        ],
        widths=[1.15, 1.85, 1.85, 1.65],
        font_size=7.3,
    )

    add_heading(doc, "7. Estimand framework", 1)
    add_p(
        doc,
        "The estimands below follow ICH E9(R1): treatment condition, population, endpoint variable, intercurrent-event handling, population-level summary and estimator are specified before analysis. Because the current repo lacks complete real CbzP IPD, confirmatory interpretation is blocked even where estimand structure is complete.",
    )
    add_table(
        doc,
        ["Endpoint", "Population", "Variable", "Intercurrent-event strategy", "Summary / estimator", "Current evidentiary status"],
        [
            ["OS", "ITT", "Time from randomization to death from any cause.", "Treatment-policy for discontinuation/subsequent therapy; censor alive subjects at earliest last-known-alive or cut-off.", "KM medians; stratified log-rank; stratified Cox HR.", "Protocol-confirmed target; current two-arm repo result non-confirmatory due synthetic CbzP."],
            ["PFS", "ITT", "Time to first valid PSA, tumour, pain progression with disease evidence, or death.", "New anti-cancer therapy before progression censors at NACTDT-1; no post-baseline assessment censors at randomization.", "KM medians; stratified log-rank; stratified Cox HR.", "Target estimand; component provenance must be repaired/validated."],
            ["PSA response", "Baseline PSA >=20 ug/L", "Confirmed >=50% PSA decline from baseline with repeat >=3 weeks.", "Missing confirmation means non-response unless death/progression handling is shell-specified.", "Response rate with exact CI; Fisher or CMH support.", "Current programming must correct eligible denominator."],
            ["ORR", "Measurable disease", "Confirmed CR/PR per RECIST v1.0.", "Missing/NE response not responder; new lesion/progression as PD.", "Response rate with exact CI; supportive duration if available.", "RECIST v1.0 required; exploratory hybrids must be labeled."],
            ["Pain response", "Pain response evaluable", "PPI or analgesic-score improvement maintained >=3 weeks.", "Insufficient diary data yields not evaluable; missing confirmation means non-response.", "Response rate and exact CI.", "Requires 5-of-7 diary evaluability implementation."],
            ["Safety", "Safety/all-treated", "TEAEs, SAEs, grade >=3 events, deaths, lab toxicities.", "Treatment discontinuation is part of safety experience; events through defined window included.", "Counts/percentages; exposure-adjusted summaries if specified.", "Current lab shift and synthetic-comparator safety outputs require fixes/labeling."],
        ],
        widths=[0.8, 0.9, 1.45, 1.55, 1.05, 1.25],
        font_size=7.1,
    )

    add_heading(doc, "8. Multiplicity and Type I error control", 1)
    add_p(
        doc,
        "For the original intended two-arm trial analysis, hierarchical testing is specified for the principal efficacy family. For the current repository implementation, p-values involving reconstructed CbzP are descriptive only and do not control Type I error for clinical claims.",
    )
    add_table(
        doc,
        ["Step", "Endpoint", "Population", "Decision rule"],
        [
            ["1", "Overall survival", "ITT", "Primary endpoint; proceed only if OS significant under approved alpha strategy."],
            ["2", "Progression-free survival", "ITT", "Test only if Step 1 is met."],
            ["3", "PSA response", "Baseline PSA >=20 ug/L", "Test only if Step 2 is met."],
            ["4", "Objective response rate", "Measurable disease", "Test only if Step 3 is met."],
            ["Exploratory", "Pain, TTPSA, TTUMOR, subgroups, exposure-response, Optimus", "Endpoint-specific", "No alpha protection; descriptive and hypothesis-generating."],
        ],
        widths=[0.8, 1.6, 1.3, 2.8],
    )

    add_heading(doc, "9. Primary efficacy analysis - overall survival", 1)
    add_heading(doc, "9.1 Definition and censoring", 2)
    add_bullets(
        doc,
        [
            "OS is time from randomization date to death from any cause.",
            "Subjects alive at analysis cut-off are censored at the earlier of last date known alive and 25-Sep-2009.",
            "Deaths after documented withdrawal but before cut-off are events if source evidence supports death date.",
            "Analysis variable: ADTTE where PARAMCD='OS'; CNSR=0 event, CNSR=1 censored; AVAL is days from randomization plus one if day-count convention requires inclusive duration.",
        ],
    )
    add_heading(doc, "9.2 Primary estimator", 2)
    add_bullets(
        doc,
        [
            "KM curves by randomized treatment.",
            "Median OS with 95% CI by Brookmeyer-Crowley method.",
            "Stratified log-rank test using randomization strata: disease measurability and ECOG PS.",
            "Stratified Cox proportional hazards model using the same strata; report HR for CbzP versus MP with 95% CI.",
            "Number at risk displayed at clinically relevant months matching the publication where feasible.",
        ],
    )
    add_heading(doc, "9.3 Sensitivity and diagnostics", 2)
    add_bullets(
        doc,
        [
            "Unstratified Cox/log-rank sensitivity.",
            "RMST at 12 months and optionally 18 months when follow-up supports interpretation.",
            "PH diagnostics using log-minus-log plots and Schoenfeld residuals where real two-arm IPD is available.",
            "Censoring distribution by arm and reason.",
        ],
    )
    add_table(
        doc,
        ["Published validation target", "Value"],
        [
            ["Median OS CbzP", "15.1 months (95% CI 14.1-16.3)"],
            ["Median OS MP", "12.7 months (95% CI 11.6-13.7)"],
            ["HR CbzP vs MP", "0.70 (95% CI 0.59-0.83), p<0.0001"],
            ["Deaths at cut-off", "513 ITT deaths, 234 CbzP and 279 MP"],
        ],
        widths=[2.0, 4.5],
    )

    add_heading(doc, "10. Secondary efficacy analyses", 1)
    add_heading(doc, "10.1 Progression-free survival", 2)
    add_bullets(
        doc,
        [
            "PFS is time from randomization to the earliest of PSA progression, tumour progression by RECIST, pain progression with supporting disease evidence, or death from any cause.",
            "The component event date is the earliest valid date among components. Component source, date and priority must be retained in ADTTE traceability variables.",
            "If no post-baseline assessment exists, censor at randomization date.",
            "If new anti-cancer therapy begins before documented progression, censor at the day before new therapy start.",
            "Primary estimator mirrors OS: KM, stratified log-rank, stratified Cox HR and medians.",
        ],
    )
    add_heading(doc, "10.2 PSA response and PSA progression", 2)
    add_bullets(
        doc,
        [
            "PSA response applies only to subjects with baseline PSA >=20 ug/L.",
            "Responder is confirmed >=50% decline from baseline, with repeat PSA at least 3 weeks later.",
            "PSA progression for non-responders is >=25% increase over nadir and absolute increase >=5 ug/L, confirmed per protocol timing.",
            "PSA progression for responders is >=50% increase over nadir; do not apply the >=5 ug/L absolute floor to responders unless a later approved amendment states otherwise.",
            "Time to PSA progression is analyzed using time-to-event methods with ITT denominator for the original trial target.",
        ],
    )
    add_heading(doc, "10.3 Tumour response and time to tumour progression", 2)
    add_bullets(
        doc,
        [
            "RECIST v1.0 governs tumour response for subjects with measurable disease.",
            "ORR is confirmed CR or PR. Confirmation requires a subsequent qualifying assessment at least 4 weeks later.",
            "SD requires sufficient duration from study entry per RECIST v1.0 conventions.",
            "New lesion or unequivocal progression yields PD.",
            "Time to tumour progression is a time-to-event endpoint distinct from composite PFS.",
        ],
    )
    add_heading(doc, "10.4 Pain endpoints", 2)
    add_bullets(
        doc,
        [
            "Pain instruments are McGill-Melzack present pain intensity (PPI) and analgesic score normalized to morphine equivalents.",
            "Visit-level PPI/analgesic summaries require at least 5 of 7 expected daily values.",
            "Pain response applies to subjects with baseline median PPI >=2 or mean analgesic score >=10.",
            "Pain response requires either >=2 point PPI reduction without analgesic-score increase, or >50% analgesic-use decrease without pain increase, maintained >=3 weeks.",
            "Pain progression uses increase from baseline/reference value, not nadir, and includes palliative radiotherapy for pain.",
        ],
    )
    add_heading(doc, "10.5 Secondary endpoint model specification matrix", 2)
    add_table(
        doc,
        ["Endpoint", "Primary analysis set", "Primary estimator", "Key sensitivity"],
        [
            ["PFS", "ITT", "Stratified Cox/log-rank and KM medians.", "Exclude assessments on/after new anti-cancer therapy; component-specific audit."],
            ["TTPSA", "ITT", "TTE methods using PSA progression event.", "Responder-status rule audit; confirmation-window sensitivity."],
            ["TTUMOR", "Measurable disease / ITT target per shell", "TTE methods using RECIST PD/new lesion.", "Assessment-window and missing imaging sensitivity."],
            ["TTPAIN", "ITT with diary evaluability", "TTE methods using pain progression event.", "5-of-7 diary rule and palliative-RT-only sensitivity."],
            ["PSA response", "Baseline PSA >=20 ug/L", "Confirmed response rate and exact CI.", "Treat missing confirmation as non-response; evaluable-only support."],
            ["ORR", "Measurable disease", "Confirmed CR/PR rate and exact CI.", "Unconfirmed response support and NE handling display."],
            ["Pain response", "Pain response evaluable", "Response rate and exact CI.", "Diary completeness and baseline pain subgroup displays."],
        ],
        widths=[0.9, 1.35, 2.0, 2.25],
        font_size=7.4,
    )

    add_heading(doc, "11. Safety analyses", 1)
    add_heading(doc, "11.1 Safety population and windows", 2)
    add_bullets(
        doc,
        [
            "Safety analyses use the all-treated population and actual treatment received.",
            "TEAE window begins at first dose and ends 30 days after last dose unless a protocol-specific SAE collection rule requires a broader window.",
            "AE severity uses NCI CTCAE v3.0; coding uses MedDRA v12.0 for publication reconciliation.",
            "The safety denominator is 371 per arm for publication targets. Current repo real patient-level safety analyses are MP-only unless complete CbzP IPD is obtained.",
        ],
    )
    add_heading(doc, "11.2 Adverse events", 2)
    add_bullets(
        doc,
        [
            "Summaries include any TEAE, serious TEAE, grade >=3 TEAE, TEAE leading to dose reduction, delay, discontinuation or death.",
            "Subject-level incidence counts each subject once per SOC/PT/grade category using worst severity.",
            "Haematological safety events are derived from laboratory grades where the publication defines them as lab-based.",
            "OCCDS episode merging may be used for event episodes, but subject-level incidence must remain reviewable and traceable.",
        ],
    )
    add_heading(doc, "11.3 Laboratory analyses", 2)
    add_bullets(
        doc,
        [
            "Laboratory analyses include baseline, scheduled cycle values, worst post-baseline grade, shift tables and nadir/recovery summaries.",
            "Lab shift programming must deduplicate to one baseline and one worst post-baseline record per subject/parameter before cross-tabulation.",
            "CTCAE grades must be traceable to source grade or independently derivable from lab value, unit and reference range.",
            "Key haematology parameters include neutrophils/ANC, leukocytes, hemoglobin/anaemia and platelets/thrombocytopenia.",
        ],
    )
    add_heading(doc, "11.4 Deaths and discontinuations", 2)
    add_bullets(
        doc,
        [
            "Deaths are summarized for ITT survival and separately for safety deaths within 30 days of last dose.",
            "Discontinuation reasons must be produced from source/ADaM records. Placeholder listings are prohibited.",
            "Any death or discontinuation listing must include USUBJID, treatment, relevant dates, reason/cause, source domain and traceability variables.",
        ],
    )

    add_heading(doc, "12. Exposure, dose modification and Project Optimus exploratory analyses", 1)
    add_bullets(
        doc,
        [
            "Exposure summaries include cycles received, treatment duration, cumulative dose, planned dose, actual dose, relative dose intensity, dose reductions and delays.",
            "RDI is actual cumulative dose divided by planned cumulative dose times 100, validated against SUPPEX anchors when available.",
            "G-CSF use is summarized by timing, cycle and prophylactic/therapeutic status where source data permits.",
            "Project Optimus analyses are retrospective and exploratory. They may structure dose/exposure-response thinking but cannot establish optimized dose for a historic completed trial.",
            "Exploratory exposure-response outputs must separate real MP-only evidence from reconstructed CbzP demonstration data.",
        ],
    )

    add_heading(doc, "13. Subgroup, sensitivity and supplementary analyses", 1)
    add_bullets(
        doc,
        [
            "Subgroup analyses are descriptive; no multiplicity adjustment is applied.",
            "OS subgroup forest plots reproduce publication-defined subgroup categories where available.",
            "Subgroups include ECOG PS, measurable disease status, pain baseline status, PSA status, prior chemotherapy regimens, age group and other publication-defined factors.",
            "Subgroup counts may not sum to 755 when the publication documents missing/unknown classifications; missing status must be displayed or footnoted.",
            "Sensitivity analyses must be specified before execution and must not replace the primary estimator unless an amendment is approved.",
        ],
    )

    add_heading(doc, "14. Data handling, missing data, partial dates and visit windows", 1)
    add_table(
        doc,
        ["Topic", "Rule"],
        [
            ["Missing efficacy outcome", "Use endpoint-specific censoring or non-responder rules. Do not impute response success."],
            ["Partial dates", "Do not invent date components in source data. Use conservative bounds for classification and sensitivity analyses where partial dates affect treatment-emergent or time-to-event status."],
            ["Week-offset dates", "AE/DS timing stored as integer weeks is reconstructed using RFSTDTC plus (week-1)*7, with +/-3.5 day uncertainty disclosed."],
            ["Duplicate records", "Define parameter-specific tie-breakers before analysis. Retain flags showing selected records."],
            ["Outliers", "Do not exclude clinical values solely because they are extreme. Flag, review and document source/data-query status."],
            ["Synthetic data", "Synthetic records must carry derivation/provenance flags and must not be mixed into real-data validation denominators."],
            ["Analysis-day convention", "ADY is date minus reference start date plus one for dates on/after reference start; negative days follow CDISC convention."],
        ],
        widths=[1.45, 5.05],
    )

    add_heading(doc, "15. CDISC implementation and metadata traceability", 1)
    datasets = rows_from_workbook("Datasets")
    var_counts = Counter(r["Dataset"] for r in rows_from_workbook("Variables"))
    ds_rows = []
    for d in datasets:
        ds_rows.append([
            d["Dataset"],
            d["Class"],
            d["Structure"],
            str(var_counts.get(d["Dataset"], 0)),
            d["Description"],
        ])
    add_table(doc, ["Dataset", "Class", "Structure", "Variable count", "Purpose"], ds_rows, widths=[0.7, 1.2, 1.25, 0.8, 2.55], font_size=7.4)
    add_bullets(
        doc,
        [
            "ADSL must contain exactly one record per subject and all population, treatment, stratification and key baseline variables needed for analysis.",
            "BDS datasets must include PARAM/PARAMCD, AVAL/AVALC as applicable, analysis dates/days, visit/window variables and analysis flags.",
            "OCCDS datasets must retain occurrence identity and derivation flags needed for incidence and episode analyses.",
            "Define-XML v2.1 with ARM must align to physical datasets and SAP output catalog with no unresolved OIDs, WhereClauses, codelists or ResultDisplay references.",
            "Every ADaM variable requires origin, predecessor, method and document reference. The current workbook Documents sheet is empty and must be remediated before release.",
            "SDTMIG version claims must match the packaged tabulation datasets. If SDTMIG 3.4 is declared, the uplifted 3.4 layer must be the packaged source.",
        ],
    )

    add_heading(doc, "16. Programming, independent QC and release gates", 1)
    add_table(
        doc,
        ["Gate", "Pass criterion"],
        [
            ["SAP approval", "Signed v4.0 or later approved before production rerun."],
            ["Dependency graph", "No unexplained orphan, dangling reference, dead production code, specified-not-produced or produced-not-specified item."],
            ["Dual programming", "SAS and R outputs match on keys, row counts, typed values, labels and hashes without lossy normalization."],
            ["Metadata/data concordance", "ADaM spec, Define-XML, physical XPT and reviewer guides agree on datasets, variables, labels, types, lengths, codelists and methods."],
            ["Conformance", "Supported CORE/Pinnacle 21/Define validation run on final delivered package; every issue dispositioned."],
            ["Logs", "Production and QC logs clean or reviewed with approved exception list."],
            ["TFL QC", "Every output has shell, SAP section, program, QC result, source data and approved status."],
            ["Package integrity", "eCTD built atomically; every payload indexed; every leaf exists; hashes match final evidence manifest."],
            ["Part 11/process", "Validated execution environment, access controls, audit trail and approval evidence available for regulated claims."],
        ],
        widths=[1.55, 4.95],
    )
    add_heading(doc, "16.1 Dual-programming reconciliation specification", 2)
    add_table(
        doc,
        ["Object", "Required comparison"],
        [
            ["ADaM datasets", "Dataset existence, row count, keys, column order, type, length, label, codelist values, typed cell values and record-level hashes."],
            ["TFL numeric results", "Independent recomputation of numerator, denominator, percentage, event count, median, CI, HR and p-value from analysis datasets."],
            ["Figures", "Independent verification of input data, at-risk counts, axis labels, treatment labels, footnotes and synthetic/non-confirmatory annotation."],
            ["Listings", "Record-level source traceability; no fabricated records; all dates and subject IDs sourced from data."],
            ["Logs", "Automated scan plus human disposition of ERROR, WARNING, invalid INPUT, uninitialized variables, merge notes and package warnings."],
            ["Hash manifest", "Every release dataset/output hash equals current physical payload and eCTD/index copy."],
        ],
        widths=[1.45, 5.05],
    )
    add_heading(doc, "16.2 Release decision classes", 2)
    add_table(
        doc,
        ["Class", "Meaning"],
        [
            ["Green", "Source, program, QC, metadata, TFL and package evidence all agree; no open critical/major issue affects the object."],
            ["Amber", "Object is scientifically usable for internal review with documented non-critical limitation; not submission-ready until limitation closed or approved."],
            ["Red", "Object is blocked due confirmed defect, missing source basis, synthetic data mislabeling, metadata drift or failed validation."],
            ["Exploratory", "Object is allowed only for hypothesis generation or pipeline demonstration; cannot support confirmatory clinical claims."],
        ],
        widths=[1.2, 5.3],
    )

    add_heading(doc, "17. TFL catalog and shell control", 1)
    add_p(
        doc,
        "TFL IDs below are the controlled SAP target catalog. Current physical outputs are implementation evidence, not proof of completion. Outputs that are not produced must either be implemented and QC'd or removed by approved amendment. Outputs produced outside this catalog must be added to the SAP or removed from release.",
    )
    tfl = sap_v3_tfl_rows()
    add_table(
        doc,
        ["ID", "Title", "Dataset(s)", "Population", "Status"],
        [[r["ID"], r["Title"], r["Dataset(s)"], r["Population"], r["Status"]] for r in tfl],
        widths=[0.75, 2.2, 1.3, 0.9, 1.35],
        font_size=6.8,
    )
    add_heading(doc, "17.1 Shell requirements for every TFL", 2)
    add_bullets(
        doc,
        [
            "Shell must state population, denominator, treatment columns, statistic rows, sorting rules, precision, missing/NE handling and footnotes.",
            "Shell must identify whether the output is confirmatory, supportive, safety, exploratory or synthetic-comparator demonstration.",
            "Shell must map to SAP section and ADaM datasets/variables before programming begins.",
            "A produced output without shell approval is not release-ready even if code runs and numbers look plausible.",
            "Any output not in the controlled catalog must be removed or added by SAP amendment before package release.",
        ],
    )

    add_heading(doc, "18. Known limitations and remediation controls", 1)
    add_callout(
        doc,
        "The limitations below are not cosmetic. They control whether the project can make clinical, regulatory or submission-readiness claims.",
        "risk",
    )
    add_table(
        doc,
        ["Finding", "Severity", "Category", "Required closure"],
        findings_rows("Critical"),
        widths=[0.75, 0.75, 1.4, 3.6],
        font_size=7.2,
    )
    add_bullets(
        doc,
        [
            "Major findings F-006 through F-025 remain release-blocking unless closed or formally risk-accepted by the appropriate owner.",
            "Synthetic comparator analyses must remain labeled non-confirmatory until complete authoritative CbzP patient-level data are available and analyzed under an approved SAP.",
            "No reviewer guide may claim clean traceability until generated from reconciled metadata and final program manifests.",
        ],
    )

    add_heading(doc, "19. References", 1)
    add_table(
        doc,
        ["Reference", "URL / repo path"],
        [
            ["FDA Study Data Technical Conformance Guide, June 2026", "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/study-data-technical-conformance-guide-technical-specifications-document"],
            ["FDA Study Data Standards Resources", "https://www.fda.gov/industry/fda-data-standards-advisory-board/study-data-standards-resources"],
            ["ICH/FDA E9(R1) estimands and sensitivity analysis", "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/e9r1-statistical-principles-clinical-trials-addendum-estimands-and-sensitivity-analysis-clinical"],
            ["FDA Project Optimus", "https://www.fda.gov/about-fda/oncology-center-excellence/project-optimus"],
            ["FDA oncology dosage optimization guidance", "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/optimizing-dosage-human-prescription-drugs-and-biological-products-treatment-oncologic-diseases"],
            ["CDISC ADaMIG v1.3", "https://www.cdisc.org/standards/foundational/adam/adamig-v1-3"],
            ["CDISC ADaM v2.1", "https://www.cdisc.org/standards/foundational/adam/adam-v2-1"],
            ["CDISC SDTMIG v3.4", "https://www.cdisc.org/standards/foundational/sdtmig/sdtmig-v3-4"],
            ["CDISC Define-XML v2.1", "https://www.cdisc.org/standards/data-exchange/define-xml/define-xml-v2-1"],
            ["CDISC Controlled Terminology release 2026-03-27", "https://www.cdisc.org/standards/terminology/controlled-terminology"],
            ["21 CFR Part 11", "https://www.ecfr.gov/current/title-21/chapter-I/subchapter-A/part-11"],
            ["TROPIC protocol", "01_raw_source/Sanofi Study Protocol Tropic.pdf"],
            ["TROPIC publication", "01_raw_source/reference_literature/de_bono_lancet_2010.pdf"],
        ],
        widths=[2.5, 4.0],
        font_size=7.2,
    )

    page_break(doc)
    add_heading(doc, "Appendix A. Source register", 1)
    add_table(doc, ["Source", "File / location", "SAP use", "Reliability / caveat"], source_register_rows(), widths=[1.55, 1.65, 1.85, 1.45], font_size=7.0)

    add_heading(doc, "Appendix B. Endpoint algorithm details", 1)
    alg_rows = [
        ["OS", "ADTTE.OS", "RANDDT, DTHDT, LSTALVDT, cut-off", "Event if death observed; otherwise censor at min(last known alive, cut-off)."],
        ["PFS", "ADTTE.PFS", "ADRS PSA/tumour/pain components, death, ADCM.NACTDT", "Event date is first valid component; censor before new therapy or at randomization/no post-baseline assessment."],
        ["TTPSA", "ADTTE.TTPSA", "PSA values in ADLB/ADRS", "Apply responder/non-responder progression rules and confirmation timing."],
        ["TTUMOR", "ADTTE.TTUMOR", "RECIST ADRS/LS", "Event on radiographic PD/new lesion; censor at last evaluable tumour assessment."],
        ["TTPAIN", "ADTTE.TTPAIN", "PN/pain diary and palliative RT CM/PR", "Event on pain progression criteria; apply diary evaluability and confirmation rules."],
        ["PSA response", "ADRS.PSARESP", "Baseline/post-baseline PSA", "Confirmed >=50% reduction from baseline, repeat >=3 weeks later; population PSA baseline >=20 ug/L."],
        ["ORR", "ADRS.OBJRESP", "RECIST target, non-target, new lesions", "Confirmed CR/PR in measurable disease population."],
        ["TEAE", "ADAE", "AE, EX, death, severity, relationship", "Treatment emergent if onset/worsening on/after first dose and within safety window."],
        ["Lab shift", "ADLB", "LB value/grade, baseline flag, ANL01FL", "One baseline and one worst post-baseline per subject/parameter before shift table."],
    ]
    add_table(doc, ["Endpoint", "ADaM target", "Key sources", "Algorithm control"], alg_rows, widths=[0.9, 1.1, 1.9, 2.6], font_size=7.2)

    add_heading(doc, "Appendix C. ADaM dataset and metadata requirements", 1)
    vrows = rows_from_workbook("Variables")
    by_ds = defaultdict(list)
    for r in vrows:
        by_ds[r["Dataset"]].append(r)
    summary_rows = []
    for ds, rows in sorted(by_ds.items()):
        mandatory = sum(1 for r in rows if r.get("Mandatory") == "Yes")
        derived = sum(1 for r in rows if "Derived" in r.get("Origin", ""))
        codelisted = sum(1 for r in rows if r.get("Codelist"))
        summary_rows.append([ds, str(len(rows)), str(mandatory), str(derived), str(codelisted), ", ".join(r["Variable"] for r in rows[:10]) + (" ..." if len(rows) > 10 else "")])
    add_table(doc, ["Dataset", "Vars", "Mandatory", "Derived", "Codelisted", "Representative variables"], summary_rows, widths=[0.8, 0.55, 0.75, 0.65, 0.75, 3.0], font_size=7.0)
    add_p(
        doc,
        "Metadata remediation rule: the Documents sheet, Predecessor metadata and Method document references must be populated before Define-XML is considered release-ready. The current workbook is a starting point, not final authority.",
    )

    add_heading(doc, "Appendix D. Full TFL catalog implementation notes", 1)
    add_table(
        doc,
        ["ID", "SAP section", "Implementation note"],
        [[r["ID"], r.get("SAP §", ""), r["Status"]] for r in tfl],
        widths=[0.8, 1.0, 4.7],
        font_size=7.0,
    )

    add_heading(doc, "Appendix E. Published validation targets", 1)
    h, r = sap_v3_table(17)
    add_table(doc, h, r, widths=[1.65, 2.0, 1.0, 1.85], font_size=6.5)

    add_heading(doc, "Appendix F. Safety, death and exposure publication targets", 1)
    add_heading(doc, "F.1 Discontinuation targets", 2)
    h, r = sap_v3_table(12)
    add_table(doc, h, r, widths=[3.0, 1.7, 1.7], font_size=7.0)
    add_heading(doc, "F.2 Adverse event targets", 2)
    h, r = sap_v3_table(13)
    add_table(doc, h, r, widths=[2.2, 1.05, 1.05, 1.05, 1.05], font_size=6.6)
    add_heading(doc, "F.3 Death targets", 2)
    h, r = sap_v3_table(14)
    add_table(doc, h, r, widths=[2.6, 1.95, 1.95], font_size=7.0)
    add_heading(doc, "F.4 Exposure targets", 2)
    h, r = sap_v3_table(15)
    add_table(doc, h, r, widths=[2.6, 1.95, 1.95], font_size=7.0)

    add_heading(doc, "Appendix G. Subgroup target register", 1)
    h, r = sap_v3_table(16)
    add_table(doc, h, r, widths=[2.1, 0.8, 1.55, 2.05], font_size=6.7)

    add_heading(doc, "Appendix H. Visit-window and protocol-correction registers", 1)
    add_heading(doc, "H.1 ADLB visit-window rules", 2)
    h, r = sap_v3_table(20)
    add_table(doc, h, r, widths=[0.65, 1.25, 0.75, 0.75, 0.75, 2.35], font_size=6.5)
    add_heading(doc, "H.2 Protocol versus publication correction register", 2)
    h, r = sap_v3_table(25)
    add_table(doc, h, r, widths=[1.4, 1.55, 1.95, 0.55, 1.05], font_size=6.4)

    add_heading(doc, "Appendix I. Full ADaM variable specification summary", 1)
    var_rows = []
    for r in vrows:
        var_rows.append([
            r.get("Dataset", ""),
            r.get("Variable", ""),
            r.get("Label", ""),
            r.get("Data Type", ""),
            r.get("Mandatory", ""),
            r.get("Codelist", ""),
            r.get("Origin", ""),
        ])
    add_table(
        doc,
        ["Dataset", "Variable", "Label", "Type", "Req", "Codelist", "Origin"],
        var_rows,
        widths=[0.65, 0.9, 2.3, 0.55, 0.45, 0.8, 0.85],
        font_size=5.8,
    )

    add_heading(doc, "Appendix J. Critical audit closure checklist", 1)
    add_table(
        doc,
        ["Finding", "Closure evidence required"],
        [
            ["F-001 eCTD integrity", "Fresh atomic sequence build; no unindexed payloads; all leaves resolve; hashes match current final outputs."],
            ["F-002 SDTM metadata/data drift", "Packaged SDTM version matches Define and standards declaration; full dataset/variable reconciliation passes."],
            ["F-003 synthetic data validity", "All synthetic-comparator outputs relabeled non-confirmatory or replaced by complete authoritative CbzP IPD analysis."],
            ["F-004 false listing", "Placeholder listing removed; discontinuation listing generated from real source/ADaM with independent QC."],
            ["F-005 submission placeholders", "Real application metadata and annotated CRF supplied or package generation blocked."],
        ],
        widths=[1.75, 4.75],
        font_size=7.6,
    )

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
